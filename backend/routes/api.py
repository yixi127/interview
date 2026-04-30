from flask import Blueprint, request, jsonify, session
from models import db
from models.user import User
from models.category import Category
from models.question import Question
import re
import io

api = Blueprint("api", __name__)


def require_admin(f):
    from functools import wraps

    @wraps(f)
    def decorated(*args, **kwargs):
        print(f"DEBUG require_admin: session.is_admin = {session.get('is_admin')}")
        print(f"DEBUG require_admin: session keys = {list(session.keys())}")
        if not session.get("is_admin"):
            return jsonify({"error": "需要管理员权限"}), 401
        return f(*args, **kwargs)

    return decorated


def parse_text_content(content):
    """解析文本格式：支持多种分隔符"""
    results = []
    lines = content.strip().split("\n")

    current_title = ""
    current_answer = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        match = re.match(r"^(\d+)[\.、]\s*(.+)$", line)
        if match:
            if current_title:
                results.append(
                    {
                        "title": current_title.strip(),
                        "answer": "\n".join(current_answer).strip(),
                    }
                )

            q_text = match.group(2).strip()
            current_title = ""
            current_answer = []

            separators = ["。", ". ", "? ", "？ "]
            for sep in separators:
                if sep in q_text:
                    parts = q_text.split(sep, 1)
                    current_title = parts[0].strip()
                    if len(parts) > 1:
                        current_answer = [parts[1].strip()]
                    break
            else:
                current_title = q_text
        elif current_title:
            current_answer.append(line)

    if current_title:
        results.append(
            {
                "title": current_title.strip(),
                "answer": "\n".join(current_answer).strip(),
            }
        )

    return results


def parse_md_content(content):
    """解析Markdown格式"""
    results = parse_text_content(content)

    alt_pattern = r"(?:^|\n)##?\s*(?:题目)?\s*(\d+)[\.、]?\s*(.+?)(?=\n##|$)"
    matches = re.findall(alt_pattern, content, re.DOTALL | re.MULTILINE)

    if matches:
        results = []
        for match in matches:
            q_text = match[1].strip()
            title = q_text
            answer = ""

            for sep in ["。", ". "]:
                if sep in q_text:
                    parts = q_text.split(sep, 1)
                    title = parts[0].strip()
                    answer = parts[1].strip() if len(parts) > 1 else ""
                    break

            results.append({"title": title, "answer": answer})

    return results


def parse_doc_content(content):
    """解析Word文档格式（支持doc和docx）"""
    import tempfile
    import os

    results = []
    if not content or len(content) < 4:
        return results

    tmp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "temp")
    os.makedirs(tmp_dir, exist_ok=True)

    tmp_path = None
    try:
        suffix = ".docx" if content[:4] == b"PK\x03\x04" else ".doc"
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=tmp_dir)
        tmp_path = tmp.name
        tmp.write(content)
        tmp.close()

        if content[:4] == b"PK\x03\x04":
            try:
                from docx import Document

                doc = Document(tmp_path)
                for para in doc.paragraphs:
                    print(f"DEBUG: {para.text[:50] if para.text else ''}")
                results = _extract_questions_from_docx(doc)
            except Exception as e:
                print(f"docx error: {e}")
                results = _parse_old_doc(tmp_path)
        else:
            results = _parse_old_doc(tmp_path)

    except Exception as e:
        print(f"parse_doc_content error: {e}")
        import traceback

        traceback.print_exc()
        results = []
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except:
                pass

    return results


def _extract_questions_from_docx(doc):
    """从docx文档提取题目"""
    results = []
    current_title = ""
    current_answer = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        match = re.match(r"^(\d+)[\.、]\s*(.+)$", text)
        if match:
            if current_title:
                results.append(
                    {
                        "title": current_title.strip(),
                        "answer": "\n".join(current_answer).strip(),
                    }
                )

            q_text = match.group(2).strip()
            current_title = ""
            current_answer = []

            for sep in ["。", ". ", "? ", "？ "]:
                if sep in q_text:
                    parts = q_text.split(sep, 1)
                    current_title = parts[0].strip()
                    if len(parts) > 1:
                        current_answer = [parts[1].strip()]
                    break
            else:
                current_title = q_text
        elif current_title:
            current_answer.append(text)

    if current_title:
        results.append(
            {
                "title": current_title.strip(),
                "answer": "\n".join(current_answer).strip(),
            }
        )

    return results


def _parse_old_doc(doc_path):
    """解析旧版.doc格式（GBK编码）"""
    results = []
    text = ""

    try:
        import win32com.client as wc

        print(f"DEBUG: Opening doc file: {doc_path}")
        word = wc.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)

        text = doc.Content.Text
        print(f"DEBUG: doc text length: {len(text)}")
        doc.Close(False)
        word.Quit()

        lines = text.split("\n")
        current_title = ""
        current_answer = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            match = re.match(r"^(\d+)[\.、]\s*(.+)$", line)
            if match:
                if current_title:
                    results.append(
                        {
                            "title": current_title.strip(),
                            "answer": "\n".join(current_answer).strip(),
                        }
                    )

                q_text = match.group(2).strip()
                current_title = ""
                current_answer = []

                for sep in ["。", ". ", "? ", "？ "]:
                    if sep in q_text:
                        parts = q_text.split(sep, 1)
                        current_title = parts[0].strip()
                        if len(parts) > 1:
                            current_answer = [parts[1].strip()]
                        break
                else:
                    current_title = q_text
            elif current_title:
                current_answer.append(line)

        if current_title:
            results.append(
                {
                    "title": current_title.strip(),
                    "answer": "\n".join(current_answer).strip(),
                }
            )

    except Exception as e:
        print(f"_parse_old_doc error: {e}")
        import traceback

        traceback.print_exc()
        results = parse_text_content(text)

    return results


parse_docx_content = parse_doc_content


@api.route("/import/parse", methods=["POST"])
@require_admin
def parse_import():
    """解析导入内容，支持文本和文件"""
    import traceback

    results = []

    print(f"DEBUG: request.files = {request.files}")
    print(f"DEBUG: request.form = {request.form}")
    print(f"DEBUG: request.content_type = {request.content_type}")
    print(f"DEBUG: request.data = {request.data[:100] if request.data else 'empty'}")

    try:
        if "file" in request.files:
            file = request.files["file"]
            filename = file.filename.lower()
            print(f"DEBUG: filename = {filename}")

            if filename.endswith(".md"):
                content = file.read().decode("utf-8")
                results = parse_md_content(content)
            elif filename.endswith((".docx", ".doc")):
                content = file.read()
                print(f"DEBUG: doc content length = {len(content)}")
                results = parse_doc_content(content)
            elif filename.endswith(".txt"):
                content = file.read().decode("utf-8")
                results = parse_text_content(content)
            else:
                return jsonify(
                    {"error": "不支持的文件格式，仅支持 txt/md/doc/docx"}
                ), 400

        else:
            data = request.json or {}
            content = data.get("content", "")
            if not content:
                return jsonify({"error": "内容不能为空"}), 400
            results = parse_text_content(content)

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"解析失败: {str(e)}"}), 500

    return jsonify(results)


@api.route("/import/save", methods=["POST"])
@require_admin
def save_imported():
    """保存导入的题目"""
    data = request.json
    questions = data.get("questions", [])
    qtype = data.get("type", "八股")
    category = data.get("category")

    if not questions:
        return jsonify({"error": "没有题目数据"}), 400
    if not category:
        return jsonify({"error": "请选择分类"}), 400

    count = 0
    for q in questions:
        if q.get("title"):
            question = Question(
                title=q["title"],
                type=qtype,
                category=category,
                importance=q.get("importance", 2),
                answer=q.get("answer", ""),
            )
            db.session.add(question)
            count += 1

    db.session.commit()
    return jsonify({"message": f"成功导入 {count} 道题目", "count": count})


@api.route("/questions", methods=["GET"])
def get_questions():
    """获取题目列表"""
    category = request.args.get("category")
    importance = request.args.get("importance")
    qtype = request.args.get("type")
    archived = request.args.get("archived")
    search = request.args.get("search", "")
    page = request.args.get("page", 1, type=int)
    per_page = request.args.get("per_page", 20, type=int)

    query = Question.query
    if archived == "1":
        query = query.filter_by(is_archived=1)
    else:
        query = query.filter_by(is_archived=0)

    if category:
        query = query.filter_by(category=category)
    if importance:
        query = query.filter_by(importance=int(importance))
    if qtype:
        query = query.filter_by(type=qtype)
    if search:
        query = query.filter(Question.title.ilike(f"%{search}%"))

    total = query.count()
    questions = (
        query.order_by(Question.importance.desc(), Question.created_at.desc())
        .offset((page - 1) * per_page)
        .limit(per_page)
        .all()
    )

    return jsonify(
        {
            "items": [q.to_dict() for q in questions],
            "total": total,
            "page": page,
            "per_page": per_page,
            "pages": (total + per_page - 1) // per_page,
        }
    )


@api.route("/questions/<int:id>", methods=["GET"])
def get_question(id):
    """获取单个题目"""
    q = Question.query.get_or_404(id)
    return jsonify(q.to_dict())


@api.route("/questions", methods=["POST"])
@require_admin
def create_question():
    """创建题目"""
    data = request.json
    title = data.get("title", "").strip()

    if not title:
        return jsonify({"error": "题目不能为空"}), 400

    category = data.get("category")
    if not category:
        return jsonify({"error": "请选择分类"}), 400

    q = Question(
        title=title,
        type=data.get("type", "八股"),
        category=category,
        importance=data.get("importance", 2),
        answer=data.get("answer", ""),
    )
    db.session.add(q)
    db.session.commit()
    return jsonify(q.to_dict())


@api.route("/questions/<int:id>", methods=["PUT"])
@require_admin
def update_question(id):
    """更新题目"""
    q = Question.query.get_or_404(id)
    data = request.json

    title = data.get("title", "").strip()
    if title:
        q.title = title
    q.type = data.get("type", q.type)
    q.category = data.get("category", q.category)
    q.importance = data.get("importance", q.importance)
    q.answer = data.get("answer", q.answer)

    db.session.commit()
    return jsonify(q.to_dict())


@api.route("/questions/<int:id>", methods=["DELETE"])
@require_admin
def delete_question(id):
    """删除题目"""
    q = Question.query.get_or_404(id)
    db.session.delete(q)
    db.session.commit()
    return jsonify({"message": "删除成功"})


@api.route("/questions/<int:id>/archive", methods=["POST"])
@require_admin
def toggle_archive(id):
    """切换归档状态"""
    q = Question.query.get_or_404(id)
    q.is_archived = 1 if q.is_archived == 0 else 0
    db.session.commit()
    return jsonify(q.to_dict())


@api.route("/categories", methods=["GET"])
def get_categories():
    """获取分类列表"""
    categories = Category.query.all()
    return jsonify([{"id": c.id, "name": c.name, "type": c.type} for c in categories])


@api.route("/categories", methods=["POST"])
@require_admin
def create_category():
    """创建分类"""
    import traceback

    print(f"DEBUG create_category: request.json = {request.json}")
    print(f"DEBUG create_category: request.data = {request.data}")
    try:
        data = request.json
        print(f"DEBUG create_category: data = {data}")
    except Exception as e:
        print(f"DEBUG create_category error: {e}")
        traceback.print_exc()
        return jsonify({"error": f"请求解析失败: {str(e)}"}), 400

    name = data.get("name", "").strip()
    category_type = data.get("type", "八股")

    print(f"DEBUG: name='{name}', type='{category_type}'")

    if not name:
        return jsonify({"error": "分类名称不能为空"}), 400

    existing = Category.query.filter_by(name=name).first()
    if existing:
        return jsonify({"error": "分类已存在"}), 400

    c = Category(name=name, type=category_type)
    db.session.add(c)
    db.session.commit()
    print(f"DEBUG: Created category {c.id} - {c.name}")
    return jsonify({"id": c.id, "name": c.name, "type": c.type})


@api.route("/categories/<int:id>", methods=["DELETE"])
@require_admin
def delete_category(id):
    """删除分类"""
    c = Category.query.get_or_404(id)
    db.session.delete(c)
    db.session.commit()
    return jsonify({"message": "删除成功"})


@api.route("/categories/<int:id>", methods=["PUT"])
@require_admin
def update_category(id):
    """更新分类"""
    c = Category.query.get_or_404(id)
    data = request.json
    name = data.get("name", "").strip()

    if not name:
        return jsonify({"error": "分类名称不能为空"}), 400

    existing = Category.query.filter(Category.name == name, Category.id != id).first()
    if existing:
        return jsonify({"error": "分类已存在"}), 400

    c.name = name
    c.type = data.get("type", c.type)
    db.session.commit()
    return jsonify({"id": c.id, "name": c.name, "type": c.type})


@api.route("/login", methods=["POST"])
def login():
    """登录"""
    data = request.json
    username = data.get("username", "").strip()
    password = data.get("password", "")

    if not username or not password:
        return jsonify({"error": "用户名和密码不能为空"}), 400

    user = User.query.filter_by(username=username).first()
    if user and user.check_password(password):
        session["user_id"] = user.id
        session["username"] = user.username
        session["is_admin"] = user.is_admin
        return jsonify({"username": user.username, "is_admin": user.is_admin})
    return jsonify({"error": "用户名或密码错误"}), 401


@api.route("/logout", methods=["POST"])
def logout():
    """退出登录"""
    session.clear()
    return jsonify({"message": "退出成功"})


@api.route("/current-user", methods=["GET"])
def current_user():
    """获取当前用户"""
    if session.get("is_admin"):
        return jsonify({"username": session.get("username"), "is_admin": True})
    return jsonify({"username": None, "is_admin": False})


@api.route("/export", methods=["GET"])
@require_admin
def export_questions():
    """导出已归档题目为Markdown"""
    from datetime import datetime

    export_type = request.args.get("type", "all")
    category = request.args.get("category")

    query = Question.query.filter_by(is_archived=1)
    if export_type and export_type != "all":
        query = query.filter_by(type=export_type)
    if category:
        query = query.filter_by(category=category)

    questions = query.order_by(
        Question.type, Question.category, Question.importance.desc()
    ).all()

    md_content = (
        f"# 面试题归档\n\n导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    )

    by_type = {}
    for q in questions:
        if q.type not in by_type:
            by_type[q.type] = {}
        if q.category not in by_type[q.type]:
            by_type[q.type][q.category] = []
        by_type[q.type][q.category].append(q)

    for qtype in ["八股", "算法"]:
        if qtype not in by_type:
            continue
        md_content += f"\n## {qtype}\n\n"

        for cat, qs in by_type[qtype].items():
            md_content += f"### {cat}\n\n"
            importance_labels = {3: "[高]", 2: "[中]", 1: "[低]"}

            for q in qs:
                imp = importance_labels.get(q.importance, "")
                md_content += f"#### {imp} {q.title}\n\n"
                if q.answer:
                    md_content += f"{q.answer}\n\n"
                md_content += "---\n\n"

    return jsonify({"content": md_content})
